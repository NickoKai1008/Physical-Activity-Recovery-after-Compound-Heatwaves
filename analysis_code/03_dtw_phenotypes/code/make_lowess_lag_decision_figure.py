from __future__ import annotations

import math
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from scipy import stats


ROOT = Path(__file__).resolve().parent
SOURCE_PACKAGE = ROOT.parent
IS_PACKAGED = ROOT.name == "code" and (ROOT.parent / "data").exists()
OUTPUT = ROOT.parent / "output" / "lowess" if IS_PACKAGED else ROOT / "deliverable"
INPUT_DATA = ROOT.parent / "data" / "results" if IS_PACKAGED else SOURCE_PACKAGE / "data"
K4_FILE = INPUT_DATA / "city_cluster_optimized_12d_ward_k4.csv"
CANONICAL_DURATION_FILE = (
    INPUT_DATA / "ward_k4_12d_group_lowess_duration_summary.csv"
)
CITY_ASSIGNMENT_FILE = (
    INPUT_DATA / "dlnm_lag_assignment_main12_vs_7day_control.csv"
)

LOWESS_FRAC = 0.55
DISPLAY_FLOOR = 0.24
BOOTSTRAP_REPLICATES = 500
BOOTSTRAP_SEED = 20260727
FDR_THRESHOLDS = {
    "strong": 3.0,
    "intermediate": 2.0,
    "overall": -math.log10(0.05),
}

COLORS_K4 = {
    1: "#B43A3A",
    2: "#E39B00",
    3: "#67ACD2",
    4: "#24558E",
}
def set_style() -> None:
    plt.rcParams.update(
        {
            "font.family": "Arial",
            "font.size": 8.2,
            "axes.linewidth": 0.8,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "xtick.major.width": 0.7,
            "ytick.major.width": 0.7,
            "xtick.major.size": 3,
            "ytick.major.size": 3,
            "figure.dpi": 160,
            "savefig.dpi": 500,
            "svg.fonttype": "none",
        }
    )


def bh_adjust(p_values: np.ndarray) -> np.ndarray:
    order = np.argsort(p_values)
    ranked = p_values[order]
    adjusted = ranked * len(ranked) / (np.arange(len(ranked)) + 1)
    adjusted = np.minimum.accumulate(adjusted[::-1])[::-1]
    output = np.empty_like(adjusted)
    output[order] = np.clip(adjusted, 0, 1)
    return output


def lagwise_fdr(city_table: pd.DataFrame, partition: str) -> pd.DataFrame:
    """Reproduce the signed Stouffer/BH-FDR evidence used by the legacy LOWESS plot."""
    core = city_table.loc[
        pd.to_numeric(city_table["is_outlier"], errors="coerce").fillna(0).eq(0)
    ].copy()
    rows: list[dict[str, float | int | str]] = []
    for cluster, group in core.groupby("cluster"):
        for lag in range(1, 13):
            beta = pd.to_numeric(
                group[f"beta_lag{lag}"], errors="coerce"
            ).to_numpy()
            p_value = pd.to_numeric(
                group[f"pval_lag{lag}"], errors="coerce"
            ).to_numpy()
            valid = np.isfinite(beta) & np.isfinite(p_value)
            beta = beta[valid]
            p_value = np.clip(p_value[valid], 1e-300, 1 - 1e-15)
            city_z = stats.norm.isf(p_value / 2) * np.sign(beta)
            stouffer_z = np.sum(city_z) / math.sqrt(len(city_z))
            stouffer_p = 2 * stats.norm.sf(abs(stouffer_z))
            rows.append(
                {
                    "partition": partition,
                    "cluster": int(cluster),
                    "lag": lag,
                    "n_core_cities": len(beta),
                    "mean_beta": float(np.mean(beta)),
                    "stouffer_z": float(stouffer_z),
                    "stouffer_p": float(stouffer_p),
                }
            )

    result = pd.DataFrame(rows)
    result["stouffer_p_fdr"] = np.nan
    for _, group in result.groupby("cluster"):
        result.loc[group.index, "stouffer_p_fdr"] = bh_adjust(
            group["stouffer_p"].to_numpy()
        )
    result["neglog10_fdr"] = -np.log10(
        np.clip(result["stouffer_p_fdr"], 1e-300, 1.0)
    )
    return result


def lowess_predict(
    x: np.ndarray,
    y: np.ndarray,
    x_new: np.ndarray,
    frac: float = LOWESS_FRAC,
) -> np.ndarray:
    """Dependency-free local-linear LOWESS matching the archived Fig3 implementation."""
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    x_new = np.asarray(x_new, dtype=float)
    finite = np.isfinite(x) & np.isfinite(y)
    x = x[finite]
    y = y[finite]
    span = max(3, int(np.ceil(frac * len(x))))
    design = np.column_stack([np.ones(len(x)), x])
    prediction = np.empty_like(x_new)
    for index, x0 in enumerate(x_new):
        distance = np.abs(x - x0)
        bandwidth = np.partition(
            distance, min(span - 1, len(x) - 1)
        )[min(span - 1, len(x) - 1)]
        if bandwidth <= 1e-12:
            weight = (distance <= 1e-12).astype(float)
        else:
            scaled = np.clip(distance / bandwidth, 0, 1)
            weight = (1 - scaled**3) ** 3
        root_weight = np.sqrt(weight)
        estimate = np.linalg.lstsq(
            design * root_weight[:, None],
            y * root_weight,
            rcond=None,
        )[0]
        prediction[index] = estimate[0] + estimate[1] * x0
    return prediction


def fit_log_evidence_lowess(lagwise: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, float | int | str]] = []
    grid = np.linspace(1, 12, 1101)
    for (partition, cluster), group in lagwise.groupby(
        ["partition", "cluster"]
    ):
        group = group.sort_values("lag")
        observed = np.clip(group["neglog10_fdr"].to_numpy(float), 1e-6, None)
        fitted_log10 = lowess_predict(
            group["lag"].to_numpy(float),
            np.log10(observed),
            grid,
        )
        fitted = 10**fitted_log10
        rows.extend(
            {
                "partition": partition,
                "cluster": int(cluster),
                "lag": float(lag),
                "lowess_neglog10_fdr": float(value),
            }
            for lag, value in zip(grid, fitted)
        )
    return pd.DataFrame(rows)


def first_post_peak_crossing(
    curve: pd.DataFrame, threshold: float
) -> float | None:
    ordered = curve.sort_values("lag")
    x = ordered["lag"].to_numpy(float)
    y = ordered["lowess_neglog10_fdr"].to_numpy(float)
    # Evidence that recovers by day 12 is classified as persistence through the horizon.
    if y[-1] >= threshold:
        return None
    peak = int(np.nanargmax(y))
    for index in range(peak + 1, len(x)):
        if y[index] < threshold:
            x0, x1 = x[index - 1], x[index]
            y0, y1 = y[index - 1], y[index]
            if abs(y1 - y0) < 1e-12:
                return float(x1)
            return float(x0 + (threshold - y0) * (x1 - x0) / (y1 - y0))
    return None


def crossing_from_arrays(
    x: np.ndarray,
    y: np.ndarray,
    threshold: float,
) -> float | None:
    if y[-1] >= threshold:
        return None
    peak = int(np.nanargmax(y))
    for index in range(peak + 1, len(x)):
        if y[index] < threshold:
            x0, x1 = x[index - 1], x[index]
            y0, y1 = y[index - 1], y[index]
            if abs(y1 - y0) < 1e-12:
                return float(x1)
            return float(x0 + (threshold - y0) * (x1 - x0) / (y1 - y0))
    return None


def bootstrap_k4_uncertainty(
    city_table: pd.DataFrame,
    n_bootstrap: int = BOOTSTRAP_REPLICATES,
    seed: int = BOOTSTRAP_SEED,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    rng = np.random.default_rng(seed)
    grid = np.linspace(1, 12, 221)
    curve_rows: list[dict[str, float | int]] = []
    endpoint_rows: list[dict[str, float | int | str]] = []
    core = city_table.loc[
        pd.to_numeric(city_table["is_outlier"], errors="coerce").fillna(0).eq(0)
    ].copy()

    for cluster, group in core.groupby("cluster"):
        bootstrap_curves: list[np.ndarray] = []
        endpoint_samples = {"strong": [], "overall": []}
        group = group.reset_index(drop=True)
        for _ in range(n_bootstrap):
            sample = group.iloc[
                rng.integers(0, len(group), size=len(group))
            ]
            p_values = []
            for lag in range(1, 13):
                beta = pd.to_numeric(
                    sample[f"beta_lag{lag}"], errors="coerce"
                ).to_numpy()
                p_value = pd.to_numeric(
                    sample[f"pval_lag{lag}"], errors="coerce"
                ).to_numpy()
                valid = np.isfinite(beta) & np.isfinite(p_value)
                beta = beta[valid]
                p_value = np.clip(p_value[valid], 1e-300, 1 - 1e-15)
                city_z = stats.norm.isf(p_value / 2) * np.sign(beta)
                cluster_z = np.sum(city_z) / math.sqrt(len(city_z))
                p_values.append(
                    max(2 * stats.norm.sf(abs(cluster_z)), 1e-300)
                )
            adjusted = bh_adjust(np.asarray(p_values))
            evidence = -np.log10(np.clip(adjusted, 1e-300, 1))
            fitted = 10 ** lowess_predict(
                np.arange(1, 13, dtype=float),
                np.log10(np.clip(evidence, 1e-6, None)),
                grid,
            )
            bootstrap_curves.append(fitted)
            for endpoint, threshold_name in (
                ("strong", "strong"),
                ("overall", "overall"),
            ):
                crossing = crossing_from_arrays(
                    grid,
                    fitted,
                    FDR_THRESHOLDS[threshold_name],
                )
                endpoint_samples[endpoint].append(
                    12.0 if crossing is None else crossing
                )

        matrix = np.vstack(bootstrap_curves)
        lower, median, upper = np.quantile(
            matrix,
            [0.025, 0.5, 0.975],
            axis=0,
        )
        curve_rows.extend(
            {
                "cluster": int(cluster),
                "lag": float(lag),
                "bootstrap_low": float(low),
                "bootstrap_median": float(mid),
                "bootstrap_high": float(high),
            }
            for lag, low, mid, high in zip(grid, lower, median, upper)
        )
        for endpoint, values in endpoint_samples.items():
            values_array = np.asarray(values, dtype=float)
            endpoint_rows.append(
                {
                    "cluster": int(cluster),
                    "endpoint": endpoint,
                    "bootstrap_low": float(np.quantile(values_array, 0.025)),
                    "bootstrap_median": float(np.quantile(values_array, 0.5)),
                    "bootstrap_high": float(np.quantile(values_array, 0.975)),
                    "persistent_at_day12_fraction": float(
                        np.mean(np.isclose(values_array, 12.0))
                    ),
                    "n_bootstrap": n_bootstrap,
                }
            )

    return pd.DataFrame(curve_rows), pd.DataFrame(endpoint_rows)


def build_duration_table(curves: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, float | int | str | None]] = []
    for cluster in (1, 2, 3, 4):
        curve = curves.loc[
            curves["partition"].eq("k4_ward")
            & curves["cluster"].eq(cluster)
        ]
        crossings = {
            name: first_post_peak_crossing(curve, threshold)
            for name, threshold in FDR_THRESHOLDS.items()
        }
        strong_days = (
            12
            if crossings["strong"] is None
            else int(round(crossings["strong"]))
        )
        overall_days = (
            12
            if crossings["overall"] is None
            else int(round(crossings["overall"]))
        )
        rows.append(
            {
                "partition": "k4_ward",
                "cluster": cluster,
                "display_label": f"C{cluster}",
                "FDR001_cross": crossings["strong"],
                "FDR01_cross": crossings["intermediate"],
                "FDR05_cross": crossings["overall"],
                "strong_days": strong_days,
                "overall_days": overall_days,
            }
        )
    return pd.DataFrame(rows)


def validate_final_windows(
    durations: pd.DataFrame, canonical: pd.DataFrame
) -> pd.DataFrame:
    final = durations.loc[durations["partition"].eq("k4_ward")].copy()
    expected = canonical.loc[canonical["Group"].between(1, 4)].copy()
    expected = expected.rename(
        columns={
            "Group": "cluster",
            "Strong days": "expected_strong_days",
            "Overall days": "expected_overall_days",
        }
    )
    audit = final.merge(
        expected[
            ["cluster", "expected_strong_days", "expected_overall_days"]
        ],
        on="cluster",
        validate="one_to_one",
    )
    audit["strong_match"] = (
        audit["strong_days"].astype(int)
        == audit["expected_strong_days"].astype(int)
    )
    audit["overall_match"] = (
        audit["overall_days"].astype(int)
        == audit["expected_overall_days"].astype(int)
    )
    if not audit[["strong_match", "overall_match"]].all().all():
        raise RuntimeError("Reproduced LOWESS windows do not match the final table.")
    return audit


def build_panel_a_source_data(
    lagwise: pd.DataFrame,
    curves: pd.DataFrame,
    bootstrap_curves: pd.DataFrame,
) -> pd.DataFrame:
    output = bootstrap_curves.copy()
    output["lowess_neglog10_fdr"] = np.nan
    for cluster, group in output.groupby("cluster"):
        fitted = curves.loc[
            curves["partition"].eq("k4_ward")
            & curves["cluster"].eq(cluster)
        ].sort_values("lag")
        output.loc[group.index, "lowess_neglog10_fdr"] = np.interp(
            group["lag"].to_numpy(float),
            fitted["lag"].to_numpy(float),
            fitted["lowess_neglog10_fdr"].to_numpy(float),
        )

    observed = lagwise.loc[
        lagwise["partition"].eq("k4_ward"),
        [
            "cluster",
            "lag",
            "n_core_cities",
            "mean_beta",
            "stouffer_z",
            "stouffer_p",
            "stouffer_p_fdr",
            "neglog10_fdr",
        ],
    ].copy()
    output["lag_key"] = output["lag"].round(6)
    observed["lag_key"] = observed["lag"].astype(float).round(6)
    output = output.merge(
        observed.drop(columns="lag"),
        on=["cluster", "lag_key"],
        how="left",
        validate="many_to_one",
    ).drop(columns="lag_key")
    return output.sort_values(["cluster", "lag"]).reset_index(drop=True)


def build_panel_b_source_data(
    city_table: pd.DataFrame,
    durations: pd.DataFrame,
    city_assignment: pd.DataFrame,
    endpoint_uncertainty: pd.DataFrame,
) -> pd.DataFrame:
    counts = cluster_counts(city_table)
    rows: list[dict[str, float | int | str]] = []
    for cluster in (1, 2, 3, 4):
        duration = durations.loc[durations["cluster"].eq(cluster)].iloc[0]
        assignment = city_assignment.loc[
            city_assignment["Group"].eq(cluster)
        ].iloc[0]
        total, core = counts[cluster]
        for endpoint, estimate_column in (
            ("strong", "strong_days"),
            ("overall", "overall_days"),
        ):
            interval = endpoint_uncertainty.loc[
                endpoint_uncertainty["cluster"].eq(cluster)
                & endpoint_uncertainty["endpoint"].eq(endpoint)
            ].iloc[0]
            rows.append(
                {
                    "cluster": cluster,
                    "n_total_cities": total,
                    "n_core_cities": core,
                    "endpoint": endpoint,
                    "estimate_days": int(duration[estimate_column]),
                    "bootstrap_low": float(interval["bootstrap_low"]),
                    "bootstrap_median": float(
                        interval["bootstrap_median"]
                    ),
                    "bootstrap_high": float(interval["bootstrap_high"]),
                    "persistent_at_day12_fraction": float(
                        interval["persistent_at_day12_fraction"]
                    ),
                    "group_mean_strong": float(
                        assignment["Group mean strong"]
                    ),
                    "group_median_strong": float(
                        assignment["Group median strong"]
                    ),
                    "group_mean_overall": float(
                        assignment["Group mean overall"]
                    ),
                    "group_median_overall": float(
                        assignment["Group median overall"]
                    ),
                    "main_lag_days": int(assignment["DLNM_lag_days_main"]),
                    "strong_sensitivity_days": int(
                        assignment["DLNM_lag_days_strong_sensitivity"]
                    ),
                    "fixed_7day_control": int(
                        assignment["DLNM_lag_days_7d_control"]
                    ),
                }
            )
    return pd.DataFrame(rows)


def build_city_assignment_source_data(
    city_table: pd.DataFrame,
    city_assignment: pd.DataFrame,
    endpoint_uncertainty: pd.DataFrame,
) -> pd.DataFrame:
    assignments = city_assignment.loc[
        city_assignment["Group"].between(1, 4),
        [
            "City",
            "Group",
            "DLNM_lag_days_main",
            "DLNM_lag_days_strong_sensitivity",
            "DLNM_lag_days_7d_control",
        ],
    ].rename(
        columns={
            "City": "city",
            "Group": "cluster",
            "DLNM_lag_days_main": "main_lag_days",
            "DLNM_lag_days_strong_sensitivity": (
                "strong_sensitivity_days"
            ),
            "DLNM_lag_days_7d_control": "fixed_7day_control",
        }
    )
    city_metadata = city_table.loc[
        :, ["city", "cluster", "is_outlier", "outlier_assigned_by"]
    ].copy()
    output = assignments.merge(
        city_metadata,
        on=["city", "cluster"],
        how="left",
        validate="one_to_one",
    )
    if output["is_outlier"].isna().any() or len(output) != 63:
        raise RuntimeError("City-level k=4 assignment merge is incomplete.")

    interval = endpoint_uncertainty.loc[
        endpoint_uncertainty["endpoint"].eq("overall"),
        [
            "cluster",
            "bootstrap_low",
            "bootstrap_median",
            "bootstrap_high",
            "persistent_at_day12_fraction",
        ],
    ].rename(
        columns={
            "bootstrap_low": "cluster_ci_low",
            "bootstrap_median": "cluster_ci_median",
            "bootstrap_high": "cluster_ci_high",
        }
    )
    output = output.merge(
        interval,
        on="cluster",
        how="left",
        validate="many_to_one",
    )

    y_positions = {1: 3.0, 2: 2.0, 3: 1.0, 4: 0.0}
    output["plot_y"] = np.nan
    for cluster, group in output.groupby("cluster"):
        ordered = group.sort_values("city")
        offsets = np.linspace(-0.22, 0.22, len(ordered))
        output.loc[ordered.index, "plot_y"] = (
            y_positions[int(cluster)] + offsets
        )
    return output.sort_values(["cluster", "city"]).reset_index(drop=True)


def cluster_counts(city_table: pd.DataFrame) -> dict[int, tuple[int, int]]:
    output: dict[int, tuple[int, int]] = {}
    for cluster, group in city_table.groupby("cluster"):
        total = len(group)
        core = int(
            pd.to_numeric(group["is_outlier"], errors="coerce")
            .fillna(0)
            .eq(0)
            .sum()
        )
        output[int(cluster)] = (total, core)
    return output


def plot_evidence_panel(
    ax: plt.Axes,
    lagwise: pd.DataFrame,
    curves: pd.DataFrame,
    durations: pd.DataFrame,
    city_table: pd.DataFrame,
    partition: str,
    colors: dict[int, str],
    labels: dict[int, str],
    title: str,
    bootstrap_curves: pd.DataFrame | None = None,
) -> None:
    subset = lagwise.loc[lagwise["partition"].eq(partition)]
    curve_subset = curves.loc[curves["partition"].eq(partition)]
    duration_subset = durations.loc[durations["partition"].eq(partition)]
    counts = cluster_counts(city_table)

    for cluster in sorted(subset["cluster"].unique()):
        points = subset.loc[subset["cluster"].eq(cluster)].sort_values("lag")
        curve = curve_subset.loc[
            curve_subset["cluster"].eq(cluster)
        ].sort_values("lag")
        duration = duration_subset.loc[
            duration_subset["cluster"].eq(cluster)
        ].iloc[0]
        color = colors[int(cluster)]
        total, core = counts[int(cluster)]
        if bootstrap_curves is not None:
            interval = bootstrap_curves.loc[
                bootstrap_curves["cluster"].eq(cluster)
            ].sort_values("lag")
            ax.fill_between(
                interval["lag"].to_numpy(float),
                np.clip(
                    interval["bootstrap_low"].to_numpy(float),
                    DISPLAY_FLOOR,
                    None,
                ),
                np.clip(
                    interval["bootstrap_high"].to_numpy(float),
                    DISPLAY_FLOOR,
                    None,
                ),
                color=color,
                alpha=0.055,
                lw=0,
                zorder=1,
            )
        ax.scatter(
            points["lag"].to_numpy(float),
            np.clip(
                points["neglog10_fdr"].to_numpy(float),
                DISPLAY_FLOOR,
                None,
            ),
            s=15,
            color=color,
            edgecolor="white",
            linewidth=0.45,
            alpha=0.84,
            zorder=4,
        )
        ax.plot(
            curve["lag"].to_numpy(float),
            np.clip(
                curve["lowess_neglog10_fdr"].to_numpy(float),
                DISPLAY_FLOOR,
                None,
            ),
            color=color,
            lw=2,
            zorder=3,
        )

    c3_like = duration_subset.loc[
        duration_subset["overall_days"].astype(int).lt(12)
    ]
    if len(c3_like) == 1:
        row = c3_like.iloc[0]
        color = colors[int(row["cluster"])]
        strong_cross = float(row["FDR001_cross"])
        overall_cross = float(row["FDR05_cross"])
        ax.axvspan(
            strong_cross,
            overall_cross,
            color=color,
            alpha=0.08,
            lw=0,
            zorder=0,
        )
        ax.axvline(
            strong_cross,
            color=color,
            lw=0.9,
            ls=(0, (2.5, 2.5)),
            alpha=0.85,
        )
        ax.axvline(
            overall_cross,
            color=color,
            lw=0.9,
            ls=(0, (4, 2.8)),
            alpha=0.85,
        )

    ax.axhline(
        FDR_THRESHOLDS["overall"],
        color="#555555",
        lw=0.9,
        ls=(0, (4, 2.8)),
        zorder=1,
    )
    ax.axhline(
        FDR_THRESHOLDS["strong"],
        color="#8C8C8C",
        lw=0.8,
        ls=(0, (1.5, 2)),
        zorder=1,
    )
    ax.text(
        12.25,
        FDR_THRESHOLDS["strong"] * 1.04,
        "FDR 0.001",
        ha="right",
        va="bottom",
        fontsize=6.4,
        color="#777777",
    )
    ax.text(
        12.25,
        FDR_THRESHOLDS["overall"] * 1.04,
        "FDR 0.05",
        ha="right",
        va="bottom",
        fontsize=6.4,
        color="#555555",
    )
    ax.set_yscale("log")
    ax.set_xlim(0.75, 12.35)
    ax.set_ylim(DISPLAY_FLOOR, 360)
    ax.set_xticks(range(1, 13))
    ax.set_yticks(
        [DISPLAY_FLOOR, FDR_THRESHOLDS["overall"], 3, 10, 100, 300]
    )
    ax.set_yticklabels(["<0.25", "1.30", "3", "10", "100", "300"])
    ax.grid(axis="x", color="#ECECEC", lw=0.45)
    ax.set_xlabel("Post-heatwave lag (days)")
    ax.set_ylabel(r"FDR evidence, $-\log_{10}(q)$")
    ax.set_title(
        title,
        loc="left",
        fontsize=9.3,
        weight="bold",
        y=1.14,
        pad=0,
    )

    clusters = sorted(subset["cluster"].unique())
    for index, cluster in enumerate(clusters):
        duration = duration_subset.loc[
            duration_subset["cluster"].eq(cluster)
        ].iloc[0]
        total, core = counts[int(cluster)]
        color = colors[int(cluster)]
        y_position = 0.86 - index * 0.19
        ax.plot(
            [1.03, 1.11],
            [y_position, y_position],
            transform=ax.transAxes,
            color=color,
            lw=2.4,
            solid_capstyle="round",
            clip_on=False,
        )
        ax.text(
            1.13,
            y_position,
            (
                f"{labels[int(cluster)]}  {total}|{core}  "
                f"{int(duration['overall_days'])} d"
            ),
            transform=ax.transAxes,
            ha="left",
            va="center",
            fontsize=6.4,
            color="#222222",
            clip_on=False,
        )


def plot_assignment_panel(
    ax: plt.Axes,
    durations: pd.DataFrame,
    city_assignment: pd.DataFrame,
    endpoint_uncertainty: pd.DataFrame,
) -> None:
    final = durations.loc[durations["partition"].eq("k4_ward")].copy()
    counts = (
        city_assignment.loc[city_assignment["Group"].between(1, 4)]
        .groupby("Group")
        .size()
        .to_dict()
    )
    y_positions = {1: 3, 2: 2, 3: 1, 4: 0}

    ax.axvspan(6.7, 7.3, color="#F1F1F1", zorder=0)
    ax.axvline(7, color="#858585", lw=0.9, ls=(0, (3, 2.5)), zorder=1)
    ax.text(
        7,
        3.56,
        "fixed 7-day sensitivity",
        ha="center",
        va="bottom",
        fontsize=6.7,
        color="#666666",
    )

    for cluster in (1, 2, 3, 4):
        row = final.loc[final["cluster"].eq(cluster)].iloc[0]
        strong = int(row["strong_days"])
        overall = int(row["overall_days"])
        y = y_positions[cluster]
        color = COLORS_K4[cluster]
        ax.hlines(y, 1, 12, color="#E7E7E7", lw=6.8, zorder=1)
        ax.hlines(y, 1, overall, color=color, lw=7.2, alpha=0.26, zorder=2)
        ax.hlines(y, 1, strong, color=color, lw=3.4, alpha=0.96, zorder=3)

        for endpoint, estimate, offset, facecolor, marker in (
            ("strong", strong, 0.13, color, "o"),
            ("overall", overall, -0.13, "white", "D"),
        ):
            uncertainty = endpoint_uncertainty.loc[
                endpoint_uncertainty["cluster"].eq(cluster)
                & endpoint_uncertainty["endpoint"].eq(endpoint)
            ].iloc[0]
            low = float(uncertainty["bootstrap_low"])
            high = float(uncertainty["bootstrap_high"])
            ax.hlines(
                y + offset,
                low,
                high,
                color=color,
                lw=1.0,
                alpha=0.65,
                zorder=4,
            )
            ax.scatter(
                estimate,
                y + offset,
                s=28,
                marker=marker,
                facecolor=facecolor,
                edgecolor=color,
                linewidth=1.1,
                zorder=5,
            )

        endpoint_label = (
            f"strong \u226512; overall \u226512"
            if strong == overall == 12
            else f"strong {strong}; overall {overall}"
        )
        summary_y = 0.84 - (cluster - 1) * 0.205
        ax.text(
            1.03,
            summary_y + 0.035,
            endpoint_label,
            transform=ax.transAxes,
            va="center",
            ha="left",
            fontsize=6.8,
            color=color,
            weight="bold",
            clip_on=False,
        )
        ax.text(
            1.03,
            summary_y - 0.035,
            (
                f"mean = median (S/O): {strong}/{overall}; "
                f"main: {overall} d"
            ),
            transform=ax.transAxes,
            va="center",
            ha="left",
            fontsize=6.1,
            color="#555555",
            clip_on=False,
        )

    ax.set_xlim(0.75, 12.35)
    ax.set_ylim(-0.58, 3.73)
    ax.set_xticks(range(1, 13))
    ax.set_yticks([3, 2, 1, 0])
    ax.set_yticklabels(
        [f"C{cluster}  (n={counts.get(cluster, 0)})" for cluster in (1, 2, 3, 4)]
    )
    ax.grid(axis="x", color="#EFEFEF", lw=0.5)
    ax.set_xlabel("Lag days retained in the main DLNM")
    ax.set_title(
        "b  Endpoint-to-model assignment",
        loc="left",
        fontsize=9.3,
        weight="bold",
        y=1.11,
        pad=0,
    )
    ax.text(
        0,
        -0.19,
        "Circle: strong endpoint (FDR 0.001). Diamond: overall endpoint (FDR 0.05). Thin line: 95% city-bootstrap interval.",
        transform=ax.transAxes,
        fontsize=6.1,
        color="#5C5C5C",
        ha="left",
        va="top",
    )
    ax.text(
        0,
        1.045,
        "LOWESS thresholds \u2192 cluster endpoints \u2192 assign to members \u2192 group median overall \u2192 main DLNM",
        transform=ax.transAxes,
        fontsize=6.4,
        color="#555555",
        ha="left",
        va="center",
        clip_on=False,
    )


def plot_city_assignment_figure(
    city_source: pd.DataFrame,
    output_base: Path,
) -> None:
    figure, ax = plt.subplots(figsize=(7.4, 4.25), facecolor="white")
    y_positions = {1: 3.0, 2: 2.0, 3: 1.0, 4: 0.0}

    ax.axvspan(6.7, 7.3, color="#F2F2F2", zorder=0)
    ax.axvline(7, color="#858585", lw=0.9, ls=(0, (3, 2.5)), zorder=1)
    ax.text(
        7,
        3.43,
        "fixed 7-day sensitivity",
        ha="center",
        va="bottom",
        fontsize=6.7,
        color="#666666",
    )

    for cluster in (1, 2, 3, 4):
        group = city_source.loc[city_source["cluster"].eq(cluster)]
        color = COLORS_K4[cluster]
        y = y_positions[cluster]
        estimate = float(group["main_lag_days"].iloc[0])
        low = float(group["cluster_ci_low"].iloc[0])
        high = float(group["cluster_ci_high"].iloc[0])
        core = group.loc[pd.to_numeric(group["is_outlier"]).eq(0)]
        outliers = group.loc[pd.to_numeric(group["is_outlier"]).eq(1)]

        ax.hlines(y, low, high, color=color, lw=1.7, alpha=0.72, zorder=3)
        ax.scatter(
            core["main_lag_days"],
            core["plot_y"],
            s=19,
            color=color,
            edgecolor="white",
            linewidth=0.45,
            alpha=0.58,
            zorder=4,
        )
        if not outliers.empty:
            ax.scatter(
                outliers["main_lag_days"],
                outliers["plot_y"],
                s=25,
                facecolor="white",
                edgecolor=color,
                linewidth=1.05,
                alpha=0.95,
                zorder=5,
            )
        ax.scatter(
            estimate,
            y,
            s=54,
            marker="D",
            facecolor="white",
            edgecolor=color,
            linewidth=1.35,
            zorder=6,
        )
        ax.text(
            12.22,
            y,
            f"{int(estimate)} d",
            ha="left",
            va="center",
            fontsize=7.2,
            color=color,
            weight="bold",
        )

    counts = city_source.groupby("cluster").size().to_dict()
    ax.set_xlim(0.75, 12.65)
    ax.set_ylim(-0.58, 3.58)
    ax.set_xticks(range(1, 13))
    ax.set_yticks([3, 2, 1, 0])
    ax.set_yticklabels(
        [f"C{cluster}  (n={counts[cluster]})" for cluster in (1, 2, 3, 4)]
    )
    ax.grid(axis="x", color="#ECECEC", lw=0.5)
    ax.set_xlabel("Assigned main DLNM lag window (days)")
    ax.set_title(
        "City assignment to phenotype-specific lag windows",
        loc="left",
        fontsize=10.2,
        weight="bold",
        pad=14,
    )
    ax.text(
        0,
        1.015,
        "City points show assigned cluster windows; diamonds and horizontal lines show cluster endpoints and 95% city-bootstrap intervals.",
        transform=ax.transAxes,
        ha="left",
        va="bottom",
        fontsize=6.7,
        color="#555555",
    )
    ax.text(
        0,
        -0.19,
        "Filled circles identify core cities; open circles identify response-shape reassignments. All points show phenotype-level window assignments.",
        transform=ax.transAxes,
        ha="left",
        va="top",
        fontsize=6.3,
        color="#5C5C5C",
    )
    figure.subplots_adjust(left=0.16, right=0.94, top=0.80, bottom=0.25)
    figure.savefig(
        output_base.with_suffix(".png"),
        dpi=500,
        facecolor="white",
    )
    figure.savefig(output_base.with_suffix(".svg"), facecolor="white")
    plt.close(figure)


def write_supporting_text(
    path: Path,
    durations: pd.DataFrame,
    endpoint_uncertainty: pd.DataFrame,
) -> None:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10)
    document.add_heading("LOWESS-guided lag-window assignment", 0)

    document.add_heading("Methods", level=1)
    document.add_paragraph(
        "City-specific PPML estimates at post-heatwave lags 1-12 were retained "
        "for the 63 cities with composite-heatwave exposure. Within each Ward-DTW "
        "cluster, signed two-sided z statistics were combined by Stouffer "
        "aggregation at each lag. The 12 cluster-level P values were adjusted "
        "by the Benjamini-Hochberg procedure. To prevent extreme early-lag "
        "evidence from dominating the duration trend, LOWESS (span 0.55) was "
        "fitted to log10[-log10(q)] and back-transformed for presentation."
    )
    document.add_paragraph(
        "The strong endpoint was the first post-peak LOWESS crossing below FDR "
        "0.001; the overall endpoint used FDR 0.05. A cluster retained 12 days "
        "when its smoothed evidence remained above the corresponding threshold "
        "at day 12; this rule classifies a transient mid-lag dip followed by recovery "
        "as persistence through the evaluated horizon. For an attenuating cluster whose terminal "
        "evidence was below threshold, the first post-peak crossing defined the "
        "endpoint. Pointwise 95% uncertainty bands and endpoint intervals were "
        "estimated from 500 city-level bootstrap resamples within each cluster."
    )
    document.add_paragraph(
        "The duration is estimated once at cluster level and then assigned to "
        "every city in that phenotype. Consequently, the reported group mean "
        "and group median summarize the assigned cluster duration. The main DLNM uses the "
        "group median overall endpoint; the group median strong endpoint is "
        "retained as a stricter sensitivity specification, alongside the fixed "
        "7-day control."
    )

    document.add_heading("Results", level=1)
    document.add_paragraph(
        "C1, C2 and C4 remained above the overall FDR threshold through day 12. "
        "C3 crossed FDR 0.001 at 5.93 days and FDR 0.05 at 7.92 days, yielding "
        "rounded endpoints of 6 and 8 days. Main DLNM windows were therefore "
        "C1=12, C2=12, C3=8 and C4=12 days; a fixed 7-day window was used for "
        "sensitivity analysis."
    )

    document.add_heading("Figure legend", level=1)
    document.add_paragraph(
        "LOWESS-guided translation of post-heatwave response phenotypes into "
        "DLNM lag windows. a, Lag-wise FDR evidence for the four Ward-DTW "
        "phenotypes. Points show observed signed Stouffer evidence after "
        "Benjamini-Hochberg correction; curves are LOWESS fits and ribbons show "
        "95% city-bootstrap intervals. Horizontal rules mark FDR 0.001 and "
        "0.05. The blue shaded interval identifies the transition between the "
        "strong and overall C3 endpoints. b, Endpoint-to-model assignment. "
        "Dark and pale bars show strong and overall retained durations. Circles "
        "and diamonds identify the respective endpoints; thin horizontal lines "
        "show 95% bootstrap intervals. The dashed line marks the fixed 7-day "
        "sensitivity window."
    )

    document.add_heading("Final lag table", level=1)
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    for cell, value in zip(
        table.rows[0].cells,
        (
            "Phenotype",
            "FDR 0.001",
            "FDR 0.05",
            "Group mean overall",
            "Group median overall",
            "Main window",
            "Bootstrap note",
        ),
    ):
        cell.text = value
    final = durations.loc[durations["partition"].eq("k4_ward")]
    for _, row in final.iterrows():
        cells = table.add_row().cells
        cluster = int(row["cluster"])
        cells[0].text = f"C{cluster}"
        cells[1].text = f"{int(row['strong_days'])} days"
        cells[2].text = f"{int(row['overall_days'])} days"
        cells[3].text = f"{int(row['overall_days'])} days"
        cells[4].text = f"{int(row['overall_days'])} days"
        cells[5].text = f"{int(row['overall_days'])} days"
        interval = endpoint_uncertainty.loc[
            endpoint_uncertainty["cluster"].eq(cluster)
            & endpoint_uncertainty["endpoint"].eq("overall")
        ].iloc[0]
        cells[6].text = (
            f"95% city-bootstrap interval "
            f"{float(interval['bootstrap_low']):.2f}-"
            f"{float(interval['bootstrap_high']):.2f} days"
        )

    document.add_picture(
        str(OUTPUT / "figure" / "fig3_k4_lowess_lag_assignment.png"),
        width=Inches(7.1),
    )
    document.add_heading("City-level assignment display", level=1)
    document.add_paragraph(
        "Each city is displayed at the lag window assigned from its DTW "
        "phenotype. These points show phenotype membership and the corresponding assigned lag window. The diamond and "
        "horizontal interval are the cluster-level overall endpoint and its "
        "95% city-bootstrap interval. Core cities are filled; outliers "
        "reassigned by response shape are open."
    )
    document.add_picture(
        str(OUTPUT / "figure" / "city_lag_assignment_by_dtw_cluster.png"),
        width=Inches(7.1),
    )
    document.save(path)


def main() -> None:
    set_style()
    for folder in ("code", "data", "results", "figure", "text"):
        (OUTPUT / folder).mkdir(parents=True, exist_ok=True)

    k4 = pd.read_csv(K4_FILE)
    canonical = pd.read_csv(CANONICAL_DURATION_FILE)
    city_assignment = pd.read_csv(CITY_ASSIGNMENT_FILE)

    lagwise = lagwise_fdr(k4, "k4_ward")
    curves = fit_log_evidence_lowess(lagwise)
    durations = build_duration_table(curves)
    bootstrap_curves, endpoint_uncertainty = bootstrap_k4_uncertainty(k4)
    validate_final_windows(durations, canonical)
    panel_a = build_panel_a_source_data(
        lagwise,
        curves,
        bootstrap_curves,
    )
    panel_b = build_panel_b_source_data(
        k4,
        durations,
        city_assignment,
        endpoint_uncertainty,
    )
    city_source = build_city_assignment_source_data(
        k4,
        city_assignment,
        endpoint_uncertainty,
    )
    panel_a.to_csv(
        OUTPUT / "results" / "panel_a_k4_lowess_evidence.csv",
        index=False,
    )
    panel_b.to_csv(
        OUTPUT / "results" / "panel_b_k4_lag_assignment.csv",
        index=False,
    )
    city_source.to_csv(
        OUTPUT / "results" / "city_lag_assignment_by_dtw_cluster.csv",
        index=False,
    )

    for source in (
        K4_FILE,
        CANONICAL_DURATION_FILE,
        CITY_ASSIGNMENT_FILE,
    ):
        destination = OUTPUT / "data" / source.name
        if source.resolve() != destination.resolve():
            shutil.copy2(source, destination)
    script_destination = OUTPUT / "code" / Path(__file__).name
    if Path(__file__).resolve() != script_destination.resolve():
        shutil.copy2(Path(__file__), script_destination)

    figure = plt.figure(figsize=(9.2, 8.0), facecolor="white")
    grid = figure.add_gridspec(
        2,
        1,
        height_ratios=(1.22, 1.0),
        left=0.10,
        right=0.73,
        bottom=0.10,
        top=0.80,
        hspace=0.44,
    )
    ax_a = figure.add_subplot(grid[0, 0])
    ax_b = figure.add_subplot(grid[1, 0], sharex=ax_a)

    plot_evidence_panel(
        ax_a,
        lagwise,
        curves,
        durations,
        k4,
        "k4_ward",
        COLORS_K4,
        {cluster: f"C{cluster}" for cluster in COLORS_K4},
        "a  k=4 phenotype evidence and thresholds",
        bootstrap_curves=bootstrap_curves,
    )
    plot_assignment_panel(
        ax_b,
        durations,
        city_assignment,
        endpoint_uncertainty,
    )
    ax_a.tick_params(axis="x", labelbottom=True)

    figure.suptitle(
        "From k=4 DTW phenotypes to phenotype-specific DLNM lag windows",
        x=0.10,
        y=0.965,
        ha="left",
        fontsize=12.8,
        weight="bold",
    )
    figure.text(
        0.10,
        0.925,
        "Cluster-level signed Stouffer evidence is FDR-adjusted across 12 lags; LOWESS terminal crossings define retained windows.",
        ha="left",
        va="top",
        fontsize=8.0,
        color="#4D4D4D",
    )
    figure.text(
        0.10,
        0.892,
        "C1/C2/C4: persistent through day 12  |  C3: 6 d strong \u2192 8 d overall",
        ha="left",
        va="top",
        fontsize=7.6,
        color=COLORS_K4[3],
        weight="bold",
    )

    figure_base = OUTPUT / "figure" / "fig3_k4_lowess_lag_assignment"
    figure.savefig(figure_base.with_suffix(".png"), dpi=500, facecolor="white")
    figure.savefig(figure_base.with_suffix(".svg"), facecolor="white")
    plt.close(figure)
    plot_city_assignment_figure(
        city_source,
        OUTPUT / "figure" / "city_lag_assignment_by_dtw_cluster",
    )

    write_supporting_text(
        OUTPUT / "text" / "Fig3_k4_LOWESS_lag_assignment_methods_results.docx",
        durations,
        endpoint_uncertainty,
    )

    readme = """# Fig3 k=4 LOWESS lag-assignment decision chain

## Read the figure from top to bottom

1. **Panel a** resolves four k=4 response phenotypes. C1, C2 and C4 remain
   above the overall FDR threshold through day 12. C3 crosses FDR 0.001 at
   5.93 days and FDR 0.05 at 7.92 days.
2. **Panel b** assigns the cluster endpoints to member cities. Main DLNM
   windows are C1=12, C2=12, C3=8 and C4=12 days. A fixed 7-day window is
   retained as sensitivity analysis.
3. **Mean and median fields** summarize the cluster-level duration assigned
   to every member city. The main model uses group median overall.

## Reproducible statistical chain

- Signed two-sided city z statistics are combined by Stouffer aggregation.
- The 12 lag-specific P values are Benjamini-Hochberg adjusted within cluster.
- LOWESS with span 0.55 is fitted to `log10[-log10(FDR p)]`.
- LOWESS uncertainty is summarized by 500 within-cluster city bootstraps.
- The first post-peak crossing below FDR 0.001 defines the strong endpoint.
- The first post-peak crossing below FDR 0.05 defines the retained endpoint.
- Evidence that recovers and remains above threshold at day 12 is classified
  as persistence through the evaluated horizon.

## Core outputs

- `figure/fig3_k4_lowess_lag_assignment.png|svg`
- `figure/city_lag_assignment_by_dtw_cluster.png|svg`
- `results/panel_a_k4_lowess_evidence.csv`
- `results/panel_b_k4_lag_assignment.csv`
- `results/city_lag_assignment_by_dtw_cluster.csv`
- `text/Fig3_k4_LOWESS_lag_assignment_methods_results.docx`

## Required inputs

- `data/city_cluster_optimized_12d_ward_k4.csv`
- `data/ward_k4_12d_group_lowess_duration_summary.csv`
- `data/dlnm_lag_assignment_main12_vs_7day_control.csv`

Run `code/make_fig3_lowess_lag_decision_chain.py` from any location. The
archived copy reads only the three files in this package's `data` directory.

LOWESS uses core cities after the established outlier screen. Figure labels
report total assigned cities followed by core-city counts, separated by `|`.
"""
    (OUTPUT / "README.md").write_text(readme, encoding="utf-8")

    manifest_rows = [
        ("code/make_fig3_lowess_lag_decision_chain.py", "analysis and figure code"),
        ("data/city_cluster_optimized_12d_ward_k4.csv", "city-level PPML lag coefficients and k=4 assignments"),
        ("data/ward_k4_12d_group_lowess_duration_summary.csv", "canonical cluster endpoint check"),
        ("data/dlnm_lag_assignment_main12_vs_7day_control.csv", "final city-to-lag assignment table"),
        ("results/panel_a_k4_lowess_evidence.csv", "panel a source data"),
        ("results/panel_b_k4_lag_assignment.csv", "panel b source data"),
        ("results/city_lag_assignment_by_dtw_cluster.csv", "city-level assigned-window display data"),
        ("figure/fig3_k4_lowess_lag_assignment.png", "500-dpi preview"),
        ("figure/fig3_k4_lowess_lag_assignment.svg", "editable vector figure"),
        ("figure/city_lag_assignment_by_dtw_cluster.png", "city assignment preview"),
        ("figure/city_lag_assignment_by_dtw_cluster.svg", "editable city assignment figure"),
        ("text/Fig3_k4_LOWESS_lag_assignment_methods_results.docx", "methods, results and legend"),
        ("README.md", "package guide"),
    ]
    pd.DataFrame(manifest_rows, columns=["relative_path", "role"]).to_csv(
        OUTPUT / "FINAL_MANIFEST.csv",
        index=False,
    )

    print("Final windows")
    print(
        durations.loc[durations["partition"].eq("k4_ward"), [
            "cluster",
            "FDR001_cross",
            "FDR05_cross",
            "strong_days",
            "overall_days",
        ]].to_string(index=False)
    )
    print(f"\nOutput: {OUTPUT}")


if __name__ == "__main__":
    main()
